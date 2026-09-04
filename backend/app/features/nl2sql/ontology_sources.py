"""Ontology 構築資料の安全な保存と決定論的な多形式抽出。"""

from __future__ import annotations

import asyncio
import contextlib
import csv
import hashlib
import importlib
import io
import json
import mimetypes
import unicodedata
import zipfile
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import uuid4

from fastapi import UploadFile

from app.clients.oci_auth import load_oci_config_without_prompt
from app.settings import Settings, get_settings

from .ontology_models import (
    OntologyEvidence,
    OntologyEvidenceLocatorKind,
    OntologySourceDocument,
    OntologySourceRole,
    QaPair,
)
from .tabular_files import (
    WORKBOOK_SUFFIXES,
    XLS_OLE_SIGNATURE,
    TabularFileReadError,
    normalize_workbook_scalar,
    read_workbook_sheets,
)

SUPPORTED_ONTOLOGY_SOURCE_SUFFIXES = frozenset(
    {".pdf", ".docx", ".txt", ".md", ".csv", ".xlsx", ".xls", ".xlsm"}
)
_CHUNK_SIZE = 1024 * 1024
_QA_QUESTION_KEYS = {"QUESTION", "質問", "TEXT", "PROMPT"}
_QA_SQL_KEYS = {"SQL", "ANSWERSQL", "回答SQL", "正解SQL"}
_ALLOWED_MEDIA_TYPES: dict[str, frozenset[str]] = {
    ".pdf": frozenset({"application/pdf"}),
    ".docx": frozenset({"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}),
    ".txt": frozenset({"text/plain"}),
    ".md": frozenset({"text/markdown", "text/plain"}),
    ".csv": frozenset({"text/csv", "application/csv", "text/plain"}),
    ".xlsx": frozenset({"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}),
    ".xls": frozenset({"application/vnd.ms-excel"}),
    ".xlsm": frozenset({"application/vnd.ms-excel.sheet.macroenabled.12"}),
}


class OntologySourceError(ValueError):
    def __init__(self, code: str, message_ja: str) -> None:
        super().__init__(message_ja)
        self.code = code
        self.message_ja = message_ja


@dataclass(frozen=True)
class ExtractedSourceChunk:
    text: str
    locator_kind: OntologyEvidenceLocatorKind
    locator: str

    def evidence(self, source: OntologySourceDocument) -> OntologyEvidence:
        normalized = normalize_source_text(self.text)
        return OntologyEvidence(
            source_document_id=source.id,
            source_sha256=source.sha256,
            locator_kind=self.locator_kind,
            locator=self.locator,
            excerpt_hash=hashlib.sha256(normalized.encode("utf-8")).hexdigest(),
            excerpt_ja=normalized[:500],
        )


@dataclass(frozen=True)
class ExtractedOntologySource:
    chunks: list[ExtractedSourceChunk]
    qa_pairs: list[QaPair]
    warnings_ja: list[str]

    @property
    def business_text(self) -> str:
        return "\n\n".join(chunk.text for chunk in self.chunks if chunk.text.strip())


def normalize_source_text(value: str) -> str:
    lines = [" ".join(line.split()) for line in unicodedata.normalize("NFC", value).splitlines()]
    return "\n".join(line for line in lines if line).strip()


def decode_source_text(content: bytes) -> str:
    """日本語資料で一般的な encoding を固定順で判定し、置換文字を混入させない。"""

    if content.startswith((b"\xff\xfe", b"\xfe\xff")):
        try:
            return content.decode("utf-16")
        except UnicodeDecodeError as exc:
            raise OntologySourceError(
                "ONTOLOGY_SOURCE_ENCODING_INVALID", "資料の文字コードを判定できません。"
            ) from exc
    for encoding in ("utf-8-sig", "cp932"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise OntologySourceError(
        "ONTOLOGY_SOURCE_ENCODING_INVALID", "資料の文字コードを判定できません。"
    )


def validate_source_signature(filename: str, content: bytes) -> None:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_ONTOLOGY_SOURCE_SUFFIXES:
        raise OntologySourceError(
            "ONTOLOGY_SOURCE_FORMAT_UNSUPPORTED",
            "PDF、DOCX、TXT、MD、CSV、XLSX、XLS、XLSM のいずれかを指定してください。",
        )
    if not content:
        raise OntologySourceError("ONTOLOGY_SOURCE_EMPTY", "空の資料は登録できません。")
    if suffix == ".pdf" and not content.startswith(b"%PDF-"):
        raise OntologySourceError("ONTOLOGY_SOURCE_MIME_MISMATCH", "PDF の内容を確認できません。")
    if suffix == ".xls" and not content.startswith(XLS_OLE_SIGNATURE):
        raise OntologySourceError(
            "ONTOLOGY_SOURCE_MIME_MISMATCH",
            "拡張子と XLS ファイル内容が一致しません。ファイルを確認して再選択してください。",
        )
    if suffix in {".docx", ".xlsx", ".xlsm"}:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = frozenset(archive.namelist())
        except (OSError, zipfile.BadZipFile) as exc:
            raise OntologySourceError(
                "ONTOLOGY_SOURCE_ARCHIVE_INVALID", "Office ファイルが壊れています。"
            ) from exc
        expected = "word/document.xml" if suffix == ".docx" else "xl/workbook.xml"
        if expected not in names:
            raise OntologySourceError(
                "ONTOLOGY_SOURCE_MIME_MISMATCH", "拡張子と Office ファイル内容が一致しません。"
            )
    if suffix in {".txt", ".md", ".csv"} and (
        b"\x00" in content[:4096]
        or content.startswith((XLS_OLE_SIGNATURE, b"PK\x03\x04", b"%PDF-"))
    ):
        raise OntologySourceError(
            "ONTOLOGY_SOURCE_BINARY_TEXT",
            "テキスト資料の拡張子と内容が一致しません。正しい形式で再選択してください。",
        )


def validate_source_path(filename: str, path: Path) -> None:
    """API のメモリへ原本全体を戻さず、保存済みファイルの magic bytes を検査する。"""

    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_ONTOLOGY_SOURCE_SUFFIXES:
        raise OntologySourceError(
            "ONTOLOGY_SOURCE_FORMAT_UNSUPPORTED",
            "PDF、DOCX、TXT、MD、CSV、XLSX、XLS、XLSM のいずれかを指定してください。",
        )
    if path.stat().st_size == 0:
        raise OntologySourceError("ONTOLOGY_SOURCE_EMPTY", "空の資料は登録できません。")
    with path.open("rb") as stream:
        prefix = stream.read(4096)
    if suffix == ".pdf" and not prefix.startswith(b"%PDF-"):
        raise OntologySourceError("ONTOLOGY_SOURCE_MIME_MISMATCH", "PDF の内容を確認できません。")
    if suffix == ".xls" and not prefix.startswith(XLS_OLE_SIGNATURE):
        raise OntologySourceError(
            "ONTOLOGY_SOURCE_MIME_MISMATCH",
            "拡張子と XLS ファイル内容が一致しません。ファイルを確認して再選択してください。",
        )
    if suffix in {".docx", ".xlsx", ".xlsm"}:
        try:
            with zipfile.ZipFile(path) as archive:
                names = frozenset(archive.namelist())
        except (OSError, zipfile.BadZipFile) as exc:
            raise OntologySourceError(
                "ONTOLOGY_SOURCE_ARCHIVE_INVALID", "Office ファイルが壊れています。"
            ) from exc
        expected = "word/document.xml" if suffix == ".docx" else "xl/workbook.xml"
        if expected not in names:
            raise OntologySourceError(
                "ONTOLOGY_SOURCE_MIME_MISMATCH", "拡張子と Office ファイル内容が一致しません。"
            )
    if suffix in {".txt", ".md", ".csv"} and (
        b"\x00" in prefix or prefix.startswith((XLS_OLE_SIGNATURE, b"PK\x03\x04", b"%PDF-"))
    ):
        raise OntologySourceError(
            "ONTOLOGY_SOURCE_BINARY_TEXT",
            "テキスト資料の拡張子と内容が一致しません。正しい形式で再選択してください。",
        )


def validate_source_media_type(filename: str, media_type: str) -> None:
    normalized = media_type.partition(";")[0].strip().lower()
    if not normalized or normalized == "application/octet-stream":
        return
    suffix = Path(filename).suffix.lower()
    if normalized not in _ALLOWED_MEDIA_TYPES.get(suffix, frozenset()):
        raise OntologySourceError(
            "ONTOLOGY_SOURCE_MIME_MISMATCH",
            "資料の Content-Type と拡張子が一致しません。",
        )


def _unlink_missing_ok(path: Path) -> None:
    path.unlink(missing_ok=True)


def _rmdir_if_empty(path: Path) -> None:
    with contextlib.suppress(OSError):
        path.rmdir()


class OntologySourceStorage:
    """local / OCI Object Storage を同じ URI 契約で扱う。"""

    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or get_settings()

    async def save_upload(
        self,
        *,
        profile_id: str,
        upload: UploadFile,
        source_role: OntologySourceRole = OntologySourceRole.SOURCE,
    ) -> OntologySourceDocument:
        filename = Path(upload.filename or "source").name
        suffix = Path(filename).suffix.lower()
        if suffix not in SUPPORTED_ONTOLOGY_SOURCE_SUFFIXES:
            raise OntologySourceError(
                "ONTOLOGY_SOURCE_FORMAT_UNSUPPORTED",
                "PDF、DOCX、TXT、MD、CSV、XLSX、XLS、XLSM のいずれかを指定してください。",
            )
        source_id = f"ontology_source_{uuid4().hex}"
        profile_storage_key = hashlib.sha256(profile_id.encode("utf-8")).hexdigest()[:24]
        target_dir = Path(self.settings.local_storage_dir).expanduser() / "ontology-sources"
        target_dir = target_dir / profile_storage_key / source_id
        await asyncio.to_thread(target_dir.mkdir, parents=True, exist_ok=True)
        target = target_dir / filename
        digest = hashlib.sha256()
        total = 0
        try:
            stream = await asyncio.to_thread(target.open, "wb")
            try:
                while chunk := await upload.read(_CHUNK_SIZE):
                    total += len(chunk)
                    if total > self.settings.max_upload_bytes:
                        raise OntologySourceError(
                            "ONTOLOGY_SOURCE_TOO_LARGE",
                            "資料がアップロード上限を超えています。",
                        )
                    digest.update(chunk)
                    await asyncio.to_thread(stream.write, chunk)
            finally:
                await asyncio.to_thread(stream.close)
        except Exception:
            await asyncio.to_thread(_unlink_missing_ok, target)
            await asyncio.to_thread(_rmdir_if_empty, target.parent)
            raise
        media_type = (
            upload.content_type or mimetypes.guess_type(filename)[0] or "application/octet-stream"
        )
        try:
            validate_source_media_type(filename, media_type)
            await asyncio.to_thread(validate_source_path, filename, target)
        except Exception:
            await asyncio.to_thread(_unlink_missing_ok, target)
            await asyncio.to_thread(_rmdir_if_empty, target.parent)
            raise
        storage_uri = str(target)
        if self.settings.upload_storage_backend.strip().lower() == "oci":
            storage_uri = await asyncio.to_thread(
                self._put_object,
                profile_storage_key,
                source_id,
                filename,
                target,
            )
            # OCI へ移した後のローカルコピーは残さない(最大 200MiB×5/build のディスクリーク防止)
            await asyncio.to_thread(_unlink_missing_ok, target)
            await asyncio.to_thread(_rmdir_if_empty, target.parent)
        return OntologySourceDocument(
            id=source_id,
            profile_id=profile_id,
            filename=filename,
            source_role=source_role,
            media_type=media_type,
            size_bytes=total,
            sha256=digest.hexdigest(),
            storage_uri=storage_uri,
        )

    def load(self, document: OntologySourceDocument) -> bytes:
        if document.storage_uri.startswith("oci://"):
            _, _, remainder = document.storage_uri.partition("oci://")
            bucket, _, object_name = remainder.partition("/")
            response = self._object_storage_client().get_object(
                self.settings.object_storage_namespace,
                bucket,
                object_name,
            )
            return bytes(response.data.content)
        return Path(document.storage_uri).read_bytes()

    def delete(self, document: OntologySourceDocument) -> None:
        """アップロード実体を削除する(Profile 削除時の cleanup 用)。失敗は握り潰さない。"""

        if document.storage_uri.startswith("oci://"):
            _, _, remainder = document.storage_uri.partition("oci://")
            bucket, _, object_name = remainder.partition("/")
            self._object_storage_client().delete_object(
                self.settings.object_storage_namespace,
                bucket,
                object_name,
            )
            return
        target = Path(document.storage_uri)
        target.unlink(missing_ok=True)
        with contextlib.suppress(OSError):
            target.parent.rmdir()

    def _put_object(
        self, profile_storage_key: str, source_id: str, filename: str, target: Path
    ) -> str:
        object_name = f"ontology-sources/{profile_storage_key}/{source_id}/{filename}"
        with target.open("rb") as stream:
            self._object_storage_client().put_object(
                self.settings.object_storage_namespace,
                self.settings.object_storage_bucket,
                object_name,
                stream,
            )
        return f"oci://{self.settings.object_storage_bucket}/{object_name}"

    def _object_storage_client(self) -> Any:
        oci_config = importlib.import_module("oci.config")
        object_storage = importlib.import_module("oci.object_storage")
        config = load_oci_config_without_prompt(
            oci_config,
            self.settings.oci_config_file,
            self.settings.resolved_oci_config_profile,
            region=self.settings.object_storage_region or self.settings.oci_region or None,
        )
        return object_storage.ObjectStorageClient(config)


def extract_ontology_source(
    source: OntologySourceDocument,
    content: bytes,
    *,
    vlm_page_runner: Callable[[bytes, int], str] | None = None,
) -> ExtractedOntologySource:
    validate_source_signature(source.filename, content)
    suffix = Path(source.filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        return _extract_text(content)
    if suffix == ".csv":
        return _extract_delimited(source.filename, content, delimiter=",")
    if suffix in WORKBOOK_SUFFIXES:
        return _extract_workbook(source.filename, content)
    if suffix == ".docx":
        return _extract_docx(content)
    if suffix == ".pdf":
        return _extract_pdf(content, vlm_page_runner=vlm_page_runner)
    raise OntologySourceError("ONTOLOGY_SOURCE_FORMAT_UNSUPPORTED", "未対応の資料形式です。")


def _extract_text(content: bytes) -> ExtractedOntologySource:
    text = normalize_source_text(decode_source_text(content))
    chunks = [
        ExtractedSourceChunk(line, OntologyEvidenceLocatorKind.LINE, f"line:{index}")
        for index, line in enumerate(text.splitlines(), start=1)
        if line
    ]
    return ExtractedOntologySource(chunks=chunks, qa_pairs=[], warnings_ja=[])


def _header_key(value: str) -> str:
    return value.strip().upper().replace("_", "")


def _qa_header_indexes(headers: list[str]) -> tuple[int, int] | None:
    normalized = [_header_key(value) for value in headers]
    question_index = next(
        (i for i, value in enumerate(normalized) if value in _QA_QUESTION_KEYS), None
    )
    sql_index = next((i for i, value in enumerate(normalized) if value in _QA_SQL_KEYS), None)
    if question_index is None or sql_index is None:
        return None
    return question_index, sql_index


def _qa_pairs_from_rows(rows: list[list[str]]) -> list[QaPair]:
    if not rows:
        return []
    qa_indexes = _qa_header_indexes(rows[0])
    if qa_indexes is None:
        return []
    question_index, sql_index = qa_indexes
    result: list[QaPair] = []
    for row in rows[1:]:
        question = row[question_index].strip() if len(row) > question_index else ""
        sql = row[sql_index].strip() if len(row) > sql_index else ""
        if question and sql and sql.split(None, 1)[0].upper() in {"SELECT", "WITH"}:
            result.append(QaPair(question=question, sql=sql))
    return result


def _column_letter(index: int) -> str:
    """0-origin の列番号を Excel と同じ A/B/.../AA 形式へ変換する。"""

    number = index + 1
    letters = ""
    while number:
        number, remainder = divmod(number - 1, 26)
        letters = chr(ord("A") + remainder) + letters
    return letters


def _tabular_headers(raw_headers: list[str]) -> list[str]:
    headers: list[str] = []
    seen_base: set[str] = set()
    seen_labels: set[str] = set()
    for index, raw_header in enumerate(raw_headers):
        column = _column_letter(index)
        base = normalize_source_text(raw_header) or f"column_{column}"
        label = base
        if base in seen_base:
            label = f"{base}__{column}"
            suffix = 2
            while label in seen_labels:
                label = f"{base}__{column}_{suffix}"
                suffix += 1
        seen_base.add(base)
        seen_labels.add(label)
        headers.append(label)
    return headers


def _row_fields(headers: list[str], row: list[str]) -> dict[str, str]:
    fields: dict[str, str] = {}
    for index, value in enumerate(row):
        normalized = normalize_source_text(value)
        if not normalized:
            continue
        header = headers[index] if index < len(headers) else f"column_{_column_letter(index)}"
        fields[header] = normalized
    return fields


def _json_line(value: dict[str, Any]) -> str:
    return json.dumps(
        {key: item for key, item in value.items() if item not in ("", None, [], {})},
        ensure_ascii=False,
        separators=(",", ":"),
    )


def _qa_jsonl_chunks_from_rows(
    rows: list[list[str]],
    *,
    source_file: str,
    source_type: str,
    sheet: str = "",
    locator_kind: OntologyEvidenceLocatorKind,
) -> list[ExtractedSourceChunk] | None:
    if not rows:
        return None
    qa_indexes = _qa_header_indexes(rows[0])
    if qa_indexes is None:
        return None
    question_index, sql_index = qa_indexes
    headers = _tabular_headers(rows[0])
    chunks: list[ExtractedSourceChunk] = []
    for row_number, row in enumerate(rows[1:], start=2):
        question = row[question_index].strip() if len(row) > question_index else ""
        sql = row[sql_index].strip() if len(row) > sql_index else ""
        if not (question and sql and sql.split(None, 1)[0].upper() in {"SELECT", "WITH"}):
            continue
        block = {
            "source_file": source_file,
            "source_type": source_type,
            "block_type": "qa_pair",
            "sheet": sheet,
            "row_start": row_number,
            "row_end": row_number,
            "question": question,
            "sql": sql,
            "fields": _row_fields(headers, row),
        }
        locator = f"sheet:{sheet};row:{row_number}" if sheet else f"row:{row_number}"
        chunks.append(ExtractedSourceChunk(_json_line(block), locator_kind, locator))
    return chunks


def _extract_delimited(
    filename: str,
    content: bytes,
    *,
    delimiter: str,
) -> ExtractedOntologySource:
    text = decode_source_text(content)
    rows = [
        [normalize_source_text(value) for value in row]
        for row in csv.reader(io.StringIO(text), delimiter=delimiter)
    ]
    qa_chunks = _qa_jsonl_chunks_from_rows(
        rows,
        source_file=filename,
        source_type="csv",
        locator_kind=OntologyEvidenceLocatorKind.QA_ROW,
    )
    if qa_chunks is not None:
        return ExtractedOntologySource(
            chunks=qa_chunks, qa_pairs=_qa_pairs_from_rows(rows), warnings_ja=[]
        )
    chunks = [
        ExtractedSourceChunk(
            json.dumps(row, ensure_ascii=False),
            OntologyEvidenceLocatorKind.QA_ROW if index > 1 else OntologyEvidenceLocatorKind.LINE,
            f"row:{index}",
        )
        for index, row in enumerate(rows, start=1)
        if any(row)
    ]
    return ExtractedOntologySource(
        chunks=chunks, qa_pairs=_qa_pairs_from_rows(rows), warnings_ja=[]
    )


def _extract_workbook(filename: str, content: bytes) -> ExtractedOntologySource:
    try:
        workbook_sheets = read_workbook_sheets(filename, content)
    except TabularFileReadError as exc:
        raise OntologySourceError("ONTOLOGY_SOURCE_EXCEL_INVALID", str(exc)) from exc
    chunks: list[ExtractedSourceChunk] = []
    qa_pairs: list[QaPair] = []
    for sheet in workbook_sheets:
        rows = [
            [normalize_source_text(normalize_workbook_scalar(value)) for value in row]
            for row in sheet.rows
        ]
        qa_pairs.extend(_qa_pairs_from_rows(rows))
        qa_chunks = _qa_jsonl_chunks_from_rows(
            rows,
            source_file=filename,
            source_type="excel",
            sheet=sheet.title,
            locator_kind=OntologyEvidenceLocatorKind.SHEET_ROW,
        )
        if qa_chunks is not None:
            chunks.extend(qa_chunks)
            continue
        chunks.extend(
            ExtractedSourceChunk(
                json.dumps(row, ensure_ascii=False),
                OntologyEvidenceLocatorKind.SHEET_ROW,
                f"sheet:{sheet.title};row:{index}",
            )
            for index, row in enumerate(rows, start=1)
            if any(row)
        )
    return ExtractedOntologySource(chunks=chunks, qa_pairs=qa_pairs, warnings_ja=[])


def _extract_docx(content: bytes) -> ExtractedOntologySource:
    from docx import Document

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise OntologySourceError(
            "ONTOLOGY_SOURCE_DOCX_INVALID", "Word 資料を読み取れません。"
        ) from exc
    chunks = [
        ExtractedSourceChunk(
            normalize_source_text(paragraph.text),
            OntologyEvidenceLocatorKind.PARAGRAPH,
            f"paragraph:{index}",
        )
        for index, paragraph in enumerate(document.paragraphs, start=1)
        if normalize_source_text(paragraph.text)
    ]
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [normalize_source_text(cell.text) for cell in row.cells]
            if any(values):
                chunks.append(
                    ExtractedSourceChunk(
                        json.dumps(values, ensure_ascii=False),
                        OntologyEvidenceLocatorKind.SHEET_ROW,
                        f"table:{table_index};row:{row_index}",
                    )
                )
    return ExtractedOntologySource(chunks=chunks, qa_pairs=[], warnings_ja=[])


def _extract_pdf(
    content: bytes,
    *,
    vlm_page_runner: Callable[[bytes, int], str] | None,
) -> ExtractedOntologySource:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            raise OntologySourceError(
                "ONTOLOGY_SOURCE_PDF_ENCRYPTED", "暗号化 PDF は使用できません。"
            )
    except OntologySourceError:
        raise
    except Exception as exc:
        raise OntologySourceError(
            "ONTOLOGY_SOURCE_PDF_INVALID", "PDF 資料を読み取れません。"
        ) from exc
    chunks: list[ExtractedSourceChunk] = []
    warnings: list[str] = []
    skipped_pages: list[str] = []
    pdfium_document: Any | None = None
    for index, page in enumerate(reader.pages, start=1):
        text = normalize_source_text(page.extract_text() or "")
        if not text and vlm_page_runner is not None:
            try:
                if pdfium_document is None:
                    import pypdfium2  # type: ignore[import-untyped]

                    pdfium_document = pypdfium2.PdfDocument(content)
                bitmap = pdfium_document[index - 1].render(scale=2)
                image = bitmap.to_pil()
                image_buffer = io.BytesIO()
                image.save(image_buffer, format="JPEG", quality=90)
                text = normalize_source_text(vlm_page_runner(image_buffer.getvalue(), index))
            except Exception:
                # 1 ページの OCR 失敗で資料全体を失敗させない(警告してスキップ)
                warnings.append(f"PDF page:{index} の VLM OCR に失敗したためスキップしました。")
                skipped_pages.append(f"page:{index}")
                continue
        if not text:
            # テキスト無しページ(画像のみ等)は警告してスキップ。全ページ空のときだけ失敗
            warnings.append(f"PDF page:{index} はテキストを抽出できないためスキップしました。")
            skipped_pages.append(f"page:{index}")
            continue
        chunks.append(ExtractedSourceChunk(text, OntologyEvidenceLocatorKind.PAGE, f"page:{index}"))
    if not chunks:
        skipped_detail = f"({', '.join(skipped_pages)})" if skipped_pages else ""
        raise OntologySourceError(
            "ONTOLOGY_SOURCE_PDF_TEXT_MISSING",
            f"PDF のどのページ{skipped_detail}からもテキストを抽出できませんでした。"
            "OCR 設定を確認してください。",
        )
    return ExtractedOntologySource(chunks=chunks, qa_pairs=[], warnings_ja=warnings)
