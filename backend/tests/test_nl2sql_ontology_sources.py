from __future__ import annotations

import hashlib
import io
import tempfile
from pathlib import Path
from typing import BinaryIO, cast

import pytest
from docx import Document
from fastapi import UploadFile
from openpyxl import Workbook  # type: ignore[import-untyped]
from pypdf import PdfWriter

from app.features.nl2sql.ontology_models import (
    OntologyEvidenceLocatorKind,
    OntologySourceDocument,
)
from app.features.nl2sql.ontology_sources import (
    OntologySourceError,
    OntologySourceStorage,
    extract_ontology_source,
    validate_source_media_type,
    validate_source_signature,
)
from app.settings import Settings


def _source(filename: str, content: bytes) -> OntologySourceDocument:
    return OntologySourceDocument(
        id=f"source-{filename}",
        profile_id="sales",
        filename=filename,
        size_bytes=len(content),
        sha256=hashlib.sha256(content).hexdigest(),
        storage_uri=f"/tmp/{filename}",
    )


def _workbook_bytes() -> bytes:
    workbook = Workbook()
    first = workbook.active
    first.title = "Q&A"
    first.append(["QUESTION", "SQL"])
    first.append(["受注件数", "SELECT COUNT(*) FROM APP.ORDERS"])
    second = workbook.create_sheet("用語")
    second.append(["受注", "注文"])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _multi_sheet_qa_workbook_bytes() -> bytes:
    workbook = Workbook()
    first = workbook.active
    first.title = "Q&A-1"
    first.append(["QUESTION", "SQL"])
    first.append(["受注件数", "SELECT COUNT(*) FROM APP.ORDERS"])
    second = workbook.create_sheet("Q&A-2")
    second.append(["QUESTION", "SQL"])
    second.append(["顧客一覧", "SELECT * FROM APP.CUSTOMERS"])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("受注は顧客に紐づきます。")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "状態"
    table.rows[0].cells[1].text = "受付中"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.parametrize(
    ("filename", "content", "locator"),
    [
        ("business.txt", "受注\n顧客".encode(), OntologyEvidenceLocatorKind.LINE),
        ("business.md", "# 受注".encode(), OntologyEvidenceLocatorKind.LINE),
        (
            "qa.csv",
            "QUESTION,SQL\n受注件数,SELECT COUNT(*) FROM APP.ORDERS\n".encode(),
            OntologyEvidenceLocatorKind.LINE,
        ),
        ("model.xlsx", _workbook_bytes(), OntologyEvidenceLocatorKind.SHEET_ROW),
        ("model.xlsm", _workbook_bytes(), OntologyEvidenceLocatorKind.SHEET_ROW),
        ("rules.docx", _docx_bytes(), OntologyEvidenceLocatorKind.PARAGRAPH),
    ],
)
def test_supported_sources_preserve_locator_and_hash(
    filename: str,
    content: bytes,
    locator: OntologyEvidenceLocatorKind,
) -> None:
    source = _source(filename, content)
    extracted = extract_ontology_source(source, content)
    assert extracted.chunks
    assert any(chunk.locator_kind == locator for chunk in extracted.chunks)
    evidence = extracted.chunks[0].evidence(source)
    assert evidence.source_sha256 == source.sha256
    assert len(evidence.excerpt_hash) == 64
    if filename.endswith((".csv", ".xlsx", ".xlsm")):
        assert extracted.qa_pairs[0].question == "受注件数"


def test_tsv_source_is_rejected() -> None:
    content = "QUESTION\tSQL\n受注件数\tSELECT COUNT(*) FROM APP.ORDERS\n".encode()

    with pytest.raises(OntologySourceError) as unsupported:
        validate_source_signature("qa.tsv", content)

    assert unsupported.value.code == "ONTOLOGY_SOURCE_FORMAT_UNSUPPORTED"

    with pytest.raises(OntologySourceError) as extract_error:
        extract_ontology_source(_source("qa.tsv", content), content)

    assert extract_error.value.code == "ONTOLOGY_SOURCE_FORMAT_UNSUPPORTED"


def test_workbook_source_collects_qa_pairs_from_all_sheets() -> None:
    content = _multi_sheet_qa_workbook_bytes()
    extracted = extract_ontology_source(_source("qa.xlsx", content), content)

    assert [pair.question for pair in extracted.qa_pairs] == ["受注件数", "顧客一覧"]


def test_scanned_pdf_uses_vlm_and_encrypted_pdf_is_rejected() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    buffer = io.BytesIO()
    writer.write(buffer)
    content = buffer.getvalue()
    extracted = extract_ontology_source(
        _source("scan.pdf", content),
        content,
        vlm_page_runner=lambda _image, page: f"{page} ページ目の受注帳票",
    )
    assert extracted.chunks[0].locator == "page:1"
    assert "受注帳票" in extracted.chunks[0].text

    encrypted_writer = PdfWriter()
    encrypted_writer.add_blank_page(width=100, height=100)
    encrypted_writer.encrypt("secret")
    encrypted = io.BytesIO()
    encrypted_writer.write(encrypted)
    with pytest.raises(OntologySourceError, match="暗号化 PDF"):
        extract_ontology_source(
            _source("encrypted.pdf", encrypted.getvalue()), encrypted.getvalue()
        )


def test_signature_spoof_is_rejected() -> None:
    with pytest.raises(OntologySourceError) as mismatch:
        validate_source_signature("spoof.pdf", b"PK fake")
    assert mismatch.value.code == "ONTOLOGY_SOURCE_MIME_MISMATCH"
    with pytest.raises(OntologySourceError) as media_mismatch:
        validate_source_media_type("rules.pdf", "text/plain")
    assert media_mismatch.value.code == "ONTOLOGY_SOURCE_MIME_MISMATCH"


def test_cp932_text_is_decoded_without_replacement_characters() -> None:
    content = "受注\n顧客".encode("cp932")
    extracted = extract_ontology_source(_source("business.txt", content), content)

    assert [chunk.text for chunk in extracted.chunks] == ["受注", "顧客"]
    assert "�" not in extracted.business_text


@pytest.mark.asyncio
async def test_streaming_upload_enforces_size_limit_and_removes_partial_file(
    tmp_path: Path,
) -> None:
    root = str(tmp_path)
    storage = OntologySourceStorage(
        Settings(
            local_storage_dir=root,
            max_upload_bytes=4,
            upload_storage_backend="local",
        )
    )
    with tempfile.SpooledTemporaryFile() as stream:
        stream.write(b"12345")
        stream.seek(0)
        upload = UploadFile(filename="large.txt", file=cast(BinaryIO, stream))
        with pytest.raises(OntologySourceError) as too_large:
            await storage.save_upload(profile_id="sales", upload=upload)
    assert too_large.value.code == "ONTOLOGY_SOURCE_TOO_LARGE"
    assert not list(Path(root).rglob("large.txt"))


@pytest.mark.asyncio
async def test_streaming_upload_keeps_untrusted_profile_id_inside_storage_root(
    tmp_path: Path,
) -> None:
    storage = OntologySourceStorage(
        Settings(
            local_storage_dir=str(tmp_path),
            max_upload_bytes=1024,
            upload_storage_backend="local",
        )
    )
    with tempfile.SpooledTemporaryFile() as stream:
        stream.write("受注".encode())
        stream.seek(0)
        upload = UploadFile(
            filename="business.txt",
            file=cast(BinaryIO, stream),
        )
        document = await storage.save_upload(profile_id="../outside", upload=upload)

    stored = Path(document.storage_uri).resolve()
    assert stored.is_relative_to(tmp_path.resolve())
    assert stored.read_text() == "受注"


def test_pdf_page_level_failures_skip_with_warning_and_all_empty_raises() -> None:
    """1 ページの抽出失敗は警告+スキップで縮退し、全ページ空のときだけ失敗する。"""

    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.add_blank_page(width=100, height=100)
    buffer = io.BytesIO()
    writer.write(buffer)
    content = buffer.getvalue()

    # page:1 は VLM 失敗、page:2 は VLM 成功 → 資料全体は成功し警告 1 件
    def flaky_runner(_image: bytes, page: int) -> str:
        if page == 1:
            raise RuntimeError("VLM timeout")
        return f"{page} ページ目の受注帳票"

    extracted = extract_ontology_source(
        _source("scan2.pdf", content), content, vlm_page_runner=flaky_runner
    )
    assert [chunk.locator for chunk in extracted.chunks] == ["page:2"]
    assert any("page:1" in warning for warning in extracted.warnings_ja)

    # VLM なしの空 PDF は全ページ空 → 失敗(スキップしたページ一覧付き)
    with pytest.raises(OntologySourceError) as exc_info:
        extract_ontology_source(_source("scan3.pdf", content), content)
    assert exc_info.value.code == "ONTOLOGY_SOURCE_PDF_TEXT_MISSING"
    assert "page:1" in exc_info.value.message_ja
    assert "OCR 設定" in exc_info.value.message_ja
